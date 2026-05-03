"""
Tezis (konferensiya) DOCX generator.

O'zbekiston standarti (Toshkent MEPhI tezislar talabnomasi):
- TNR 14, satr oraligi 1.0
- Hoshiyalar: chap 3cm, o'ng 1.5cm, yuqori/pastki 2cm
- 1-qator chekinishi 1cm
- Maksimum 2 bet
- TITUL VARAGI YO'Q — birinchi sahifadan boshlanadi:
  1) MA'RUZA NOMI (BOSH HARFLAR, bold, markazda)
  2) Muallif F.I.O., daraja, tashkilot, shahar
  3) Email
  4) (bo'sh qator)
  5) Asosiy matn (justify, 1cm chekinishi)
  6) Adabiyotlar — [1], [2], ...
"""

from __future__ import annotations

import os
import re
import logging
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ThesisGenerator:
    """Konferensiya tezisi DOCX yaratish (1-2 bet, ramkasiz)."""

    FONT_NAME = 'Times New Roman'

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kerak")

    def create_thesis(self, content: Dict, output_path: str) -> bool:
        """
        Tezis DOCX yaratish.

        content sxemasi:
            title:        ma'ruza nomi
            authors:      [{'name': '...', 'rank': 'PhD', 'institution': '...', 'city': 'Toshkent', 'country': "O'zbekiston"}]
            email:        muallif emaili
            sections:     [{'heading': 'Kirish', 'content': '...'}, ...]  # ixtiyoriy bo'limlar
            body:         oddiy paragraf matni (sections o'rniga ham ishlatish mumkin)
            references:   ['1. ...', '2. ...']
            keywords:     ['so'z1', 'so'z2', ...]  # ixtiyoriy
        """
        try:
            doc = Document()
            self._setup_page(doc)
            self._setup_styles(doc)

            self._add_title(doc, content.get('title', 'TEZIS'))
            self._add_authors(doc, content.get('authors', []))

            email = content.get('email', '')
            if email:
                self._add_email(doc, email)

            self._add_empty_paragraph(doc)

            keywords = content.get('keywords', [])
            if keywords:
                self._add_keywords(doc, keywords)

            sections = content.get('sections', [])
            if sections:
                for sec in sections:
                    heading = sec.get('heading', '')
                    text = sec.get('content', '')
                    if heading:
                        self._add_section_heading(doc, heading)
                    self._add_body_text(doc, text)
            else:
                self._add_body_text(doc, content.get('body', ''))

            references = content.get('references', [])
            if references:
                self._add_references(doc, references)

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logger.info(f"Tezis DOCX saqlandi: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Tezis DOCX yaratishda xato: {e}", exc_info=True)
            return False

    def _setup_page(self, doc: Document):
        section = doc.sections[0]
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    def _setup_styles(self, doc: Document):
        style = doc.styles['Normal']
        style.font.name = self.FONT_NAME
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.line_spacing = 1.0
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.first_line_indent = Cm(1)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _add_title(self, doc: Document, title: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = self.FONT_NAME

    def _add_authors(self, doc: Document, authors: List[Dict]):
        if not authors:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        parts = []
        for a in authors:
            name = a.get('name', '')
            rank = a.get('rank', '')
            inst = a.get('institution', '')
            city = a.get('city', '')
            country = a.get('country', '')
            tail = ', '.join(x for x in [rank, inst, city, country] if x)
            full = f"{name} ({tail})" if tail else name
            parts.append(full)
        run = p.add_run('; '.join(parts))
        run.font.size = Pt(13)
        run.font.name = self.FONT_NAME
        run.italic = True

    def _add_email(self, doc: Document, email: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(f"E-mail: {email}")
        run.font.size = Pt(12)
        run.font.name = self.FONT_NAME

    def _add_empty_paragraph(self, doc: Document):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(0)

    def _add_keywords(self, doc: Document, keywords: List[str]):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1)
        p.paragraph_format.space_after = Pt(6)
        bold_run = p.add_run("Kalit so'zlar: ")
        bold_run.bold = True
        bold_run.font.size = Pt(13)
        bold_run.font.name = self.FONT_NAME
        text_run = p.add_run(', '.join(keywords) + '.')
        text_run.font.size = Pt(13)
        text_run.font.name = self.FONT_NAME

    def _add_section_heading(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = self.FONT_NAME

    def _add_body_text(self, doc: Document, text: str):
        if not text:
            return
        # Paragraflarga ajratish (bo'sh qator bo'yicha)
        for para in re.split(r'\n\s*\n', text.strip()):
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            self._add_runs_with_bold(p, para.replace('\n', ' '))

    def _add_runs_with_bold(self, paragraph, text: str):
        """**bold** belgilarini ajratib qo'shish."""
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                run = paragraph.add_run(part)
            run.font.size = Pt(14)
            run.font.name = self.FONT_NAME

    def _add_references(self, doc: Document, references: List[str]):
        # Adabiyotlar sarlavhasi
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(1)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run('Adabiyotlar:')
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = self.FONT_NAME

        for i, ref in enumerate(references, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            ref_text = ref if re.match(r'^\d+\.\s', ref) else f"{i}. {ref}"
            run = p.add_run(ref_text)
            run.font.size = Pt(12)
            run.font.name = self.FONT_NAME
