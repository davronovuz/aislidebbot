"""
Krossvord generator — AI'dan so'z+ta'rif olib, grid'ga joylab DOCX chiqaradi.

Algoritm (greedy placement):
1. So'zlarni eng uzunidan boshlab tartiblaymiz.
2. Birinchi so'zni grid markaziga gorizontal qo'yamiz.
3. Har keyingi so'z uchun: mavjud grid'dagi har harf bilan kesishish nuqtasini topamiz.
   To'g'ri kesishuv (bir xil harf, atrof bo'sh, ikkita parallel so'z bir-biriga tegib turmaydi) bo'lsa joylashtiramiz.
4. Joylashtirilgan so'zlarni gorizontal/vertikal raqamlab chiqamiz (NYT stilida — har so'z boshi raqamlanadi).

Output DOCX strukturasi:
  1-bet: Mavzu + bo'sh grid (raqamli hujayralar bilan, qora hujayralar fon=qora)
  2-bet: Eniga (Across) savollar + Bo'yiga (Down) savollar
  3-bet: Javoblar (to'ldirilgan grid)
"""

from __future__ import annotations

import os
import re
import json
import logging
import random
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


GRID_SIZE = 15  # 15x15 standart


class CrosswordGrid:
    """Krossvord grid algoritmi (greedy)."""

    def __init__(self, size: int = GRID_SIZE):
        self.size = size
        # None = bo'sh hujayra (krossvordda qora bo'lib chiziladi)
        # str = harf
        self.grid: List[List[Optional[str]]] = [[None] * size for _ in range(size)]
        # Joylashtirilgan so'zlar: {'word', 'clue', 'row', 'col', 'direction', 'number'}
        self.placed: List[Dict] = []

    def can_place(self, word: str, row: int, col: int, direction: str) -> int:
        """
        Tekshirish: shu pozitsiyaga so'z to'g'ri keladimi?
        Qaytaradi: kesishuv soni (>0 = bo'lishi mumkin va qancha yaxshi),
                   -1 = bo'lmaydi.
        Birinchi so'z uchun (placed bo'sh) — har qanday joy bo'ladi.
        """
        word = word.upper()
        n = len(word)

        if direction == 'across':
            if col < 0 or col + n > self.size or row < 0 or row >= self.size:
                return -1
            # Yon hujayralar (so'z oldi va keyin)
            if col > 0 and self.grid[row][col - 1] is not None:
                return -1
            if col + n < self.size and self.grid[row][col + n] is not None:
                return -1
            intersections = 0
            for i, ch in enumerate(word):
                cell = self.grid[row][col + i]
                if cell is None:
                    # Bo'sh hujayra — yuqori va pastki qo'shni harf bo'lmasligi kerak
                    # (aks holda yangi so'z hosil qiladi)
                    if row > 0 and self.grid[row - 1][col + i] is not None:
                        return -1
                    if row + 1 < self.size and self.grid[row + 1][col + i] is not None:
                        return -1
                elif cell == ch:
                    intersections += 1
                else:
                    return -1
            return intersections

        else:  # down
            if row < 0 or row + n > self.size or col < 0 or col >= self.size:
                return -1
            if row > 0 and self.grid[row - 1][col] is not None:
                return -1
            if row + n < self.size and self.grid[row + n][col] is not None:
                return -1
            intersections = 0
            for i, ch in enumerate(word):
                cell = self.grid[row + i][col]
                if cell is None:
                    if col > 0 and self.grid[row + i][col - 1] is not None:
                        return -1
                    if col + 1 < self.size and self.grid[row + i][col + 1] is not None:
                        return -1
                elif cell == ch:
                    intersections += 1
                else:
                    return -1
            return intersections

    def place(self, word: str, clue: str, row: int, col: int, direction: str):
        word = word.upper()
        for i, ch in enumerate(word):
            r = row + (i if direction == 'down' else 0)
            c = col + (i if direction == 'across' else 0)
            self.grid[r][c] = ch
        self.placed.append({
            'word': word,
            'clue': clue,
            'row': row,
            'col': col,
            'direction': direction,
        })

    def try_add(self, word: str, clue: str) -> bool:
        """Yangi so'zni eng yaxshi joyga qo'yishga urinish."""
        word = word.upper()
        if not self.placed:
            # Birinchi so'z — markazga gorizontal
            row = self.size // 2
            col = (self.size - len(word)) // 2
            if self.can_place(word, row, col, 'across') >= 0:
                self.place(word, clue, row, col, 'across')
                return True
            return False

        best = None  # (intersections, row, col, direction)
        for placed_w in self.placed:
            pw = placed_w['word']
            for i, ch in enumerate(word):
                for j, pch in enumerate(pw):
                    if ch != pch:
                        continue
                    if placed_w['direction'] == 'across':
                        # Yangi so'z VERTIKAL bo'lishi kerak (kesishuv)
                        new_row = placed_w['row'] - i
                        new_col = placed_w['col'] + j
                        score = self.can_place(word, new_row, new_col, 'down')
                        if score > 0:
                            if best is None or score > best[0]:
                                best = (score, new_row, new_col, 'down')
                    else:
                        # Yangi so'z GORIZONTAL bo'lishi kerak
                        new_row = placed_w['row'] + j
                        new_col = placed_w['col'] - i
                        score = self.can_place(word, new_row, new_col, 'across')
                        if score > 0:
                            if best is None or score > best[0]:
                                best = (score, new_row, new_col, 'across')

        if best is None:
            return False

        _, r, c, d = best
        self.place(word, clue, r, c, d)
        return True

    def number_cells(self) -> Dict[Tuple[int, int], int]:
        """
        Har so'zning birinchi hujayrasiga raqam beradi (NYT pattern).
        Hujayra raqamlanadi agar:
          - across so'z boshlansa (chap qo'shni bo'sh va o'ng tomonda harf bor)
          - down so'z boshlansa (yuqori qo'shni bo'sh va pastki tomonda harf bor)
        """
        numbers: Dict[Tuple[int, int], int] = {}
        counter = 0
        for r in range(self.size):
            for c in range(self.size):
                if self.grid[r][c] is None:
                    continue
                left_empty = (c == 0 or self.grid[r][c - 1] is None)
                top_empty = (r == 0 or self.grid[r - 1][c] is None)
                right_filled = (c + 1 < self.size and self.grid[r][c + 1] is not None)
                bottom_filled = (r + 1 < self.size and self.grid[r + 1][c] is not None)
                starts_across = left_empty and right_filled
                starts_down = top_empty and bottom_filled
                if starts_across or starts_down:
                    counter += 1
                    numbers[(r, c)] = counter
        return numbers

    def trim_bounds(self) -> Tuple[int, int, int, int]:
        """Bo'sh chetlarni kesib, ishlatilgan oraliqni qaytaradi."""
        rows = [r for r in range(self.size) if any(c is not None for c in self.grid[r])]
        cols = [c for c in range(self.size) if any(self.grid[r][c] is not None for r in range(self.size))]
        if not rows or not cols:
            return 0, 0, self.size - 1, self.size - 1
        return min(rows), min(cols), max(rows), max(cols)


class CrosswordDocxRenderer:
    """Krossvord grid'ini DOCX hujjat sifatida chizadi."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kerak")

    def create_crossword(self, grid: CrosswordGrid, topic: str, output_path: str) -> bool:
        """3 betli krossvord DOCX:
            1) topic + bo'sh grid (raqamlar + qora hujayralar)
            2) Eniga / Bo'yiga savollar
            3) Javoblar (to'ldirilgan grid)
        """
        try:
            doc = Document()
            self._setup_page(doc)

            numbers = grid.number_cells()
            r0, c0, r1, c1 = grid.trim_bounds()

            # ─── 1-bet: mavzu + bo'sh grid
            self._add_title(doc, topic)
            self._add_subtitle(doc, "Quyidagi krossvordni yeching")
            self._render_grid_table(doc, grid, numbers, r0, c0, r1, c1, fill_letters=False)
            doc.add_page_break()

            # ─── 2-bet: savollar
            self._add_clues_page(doc, grid, numbers)
            doc.add_page_break()

            # ─── 3-bet: javoblar
            self._add_title(doc, "Javoblar")
            self._add_subtitle(doc, f"Mavzu: {topic}")
            self._render_grid_table(doc, grid, numbers, r0, c0, r1, c1, fill_letters=True)

            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logger.info(f"Krossvord DOCX saqlandi: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Krossvord DOCX xato: {e}", exc_info=True)
            return False

    def _setup_page(self, doc):
        for section in doc.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)

    def _add_title(self, doc, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(20)
        run.font.name = 'Times New Roman'

    def _add_subtitle(self, doc, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    def _render_grid_table(self, doc, grid: CrosswordGrid, numbers: Dict[Tuple[int, int], int],
                           r0: int, c0: int, r1: int, c1: int, fill_letters: bool):
        rows_count = r1 - r0 + 1
        cols_count = c1 - c0 + 1
        table = doc.add_table(rows=rows_count, cols=cols_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Hujayra o'lchamini fix qilish (kvadrat — taxminan 0.7 cm)
        cell_size = Cm(0.7)
        for row_idx, table_row in enumerate(table.rows):
            table_row.height = cell_size
            for col_idx, cell in enumerate(table_row.cells):
                self._set_cell_size(cell, cell_size)
                gr = r0 + row_idx
                gc = c0 + col_idx
                letter = grid.grid[gr][gc]
                self._set_cell_borders(cell)
                if letter is None:
                    # Qora hujayra
                    self._set_cell_shading(cell, '000000')
                    cell.text = ''
                else:
                    num = numbers.get((gr, gc))
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    cell.text = ''
                    if num is not None:
                        # Raqam (kichik, yuqori chap)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
                        run = p.add_run(str(num))
                        run.font.size = Pt(7)
                        run.font.name = 'Times New Roman'
                        if fill_letters:
                            # Raqamdan keyin yangi qator + harf
                            run2 = p.add_run('\n')
                            letter_run = p.add_run(letter)
                            letter_run.bold = True
                            letter_run.font.size = Pt(14)
                            letter_run.font.name = 'Times New Roman'
                    elif fill_letters:
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(letter)
                        run.bold = True
                        run.font.size = Pt(14)
                        run.font.name = 'Times New Roman'

    def _set_cell_size(self, cell, size):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(size.emu / 635)}" w:type="dxa"/>')
        # Eski tcW bo'lsa olib tashlaymiz
        for old in tcPr.findall(qn('w:tcW')):
            tcPr.remove(old)
        tcPr.append(tcW)

    def _set_cell_shading(self, cell, hex_color: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:shd')):
            tcPr.remove(old)
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def _set_cell_borders(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(old)
        borders_xml = (
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="6" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="6" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="6" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="6" w:color="000000"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(parse_xml(borders_xml))

    def _add_clues_page(self, doc, grid: CrosswordGrid, numbers: Dict[Tuple[int, int], int]):
        # So'zlarni "across" / "down" bo'yicha guruhlash, raqamlash bo'yicha tartiblash
        across = []
        down = []
        for w in grid.placed:
            num = numbers.get((w['row'], w['col']))
            if num is None:
                continue
            entry = (num, w['clue'], w['word'])
            if w['direction'] == 'across':
                across.append(entry)
            else:
                down.append(entry)
        across.sort()
        down.sort()

        # ENIGA
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run('Eniga (Gorizontal):')
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        for num, clue, _ in across:
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent = Cm(0.5)
            cp.paragraph_format.first_line_indent = Cm(-0.5)
            cp.paragraph_format.line_spacing = 1.15
            cp.paragraph_format.space_after = Pt(2)
            num_run = cp.add_run(f"{num}. ")
            num_run.bold = True
            num_run.font.size = Pt(12)
            num_run.font.name = 'Times New Roman'
            text_run = cp.add_run(clue)
            text_run.font.size = Pt(12)
            text_run.font.name = 'Times New Roman'

        # BO'YIGA
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run("Bo'yiga (Vertikal):")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'

        for num, clue, _ in down:
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent = Cm(0.5)
            cp.paragraph_format.first_line_indent = Cm(-0.5)
            cp.paragraph_format.line_spacing = 1.15
            cp.paragraph_format.space_after = Pt(2)
            num_run = cp.add_run(f"{num}. ")
            num_run.bold = True
            num_run.font.size = Pt(12)
            num_run.font.name = 'Times New Roman'
            text_run = cp.add_run(clue)
            text_run.font.size = Pt(12)
            text_run.font.name = 'Times New Roman'


# ───────────────────────────────────────────────────────────────
# AI orqali so'z+ta'rif olish
# ───────────────────────────────────────────────────────────────

class CrosswordAIGenerator:
    """OpenAI orqali mavzuga oid 15-25 ta so'z+ta'rif juftligi olish."""

    def __init__(self, openai_key: str):
        from openai import AsyncOpenAI
        self.client = AsyncOpenAI(api_key=openai_key)

    async def generate_words(self, topic: str, count: int = 18, language: str = 'uz',
                              attempt: int = 1) -> List[Dict]:
        """
        Qaytaradi: [{'word': 'PYTHON', 'clue': 'Mashhur dasturlash tili'}, ...]
        AI dan JSON oladi, tozalaydi, kam bo'lsa qayta so'raydi (max 2 marta).
        """
        lang_label = {
            'uz': "lotin alifbosida o'zbek tilida (APOSTROFSIZ! 'so'z' o'rniga 'soz', 'o'qituvchi' o'rniga 'oqituvchi')",
            'ru': 'на русском языке (только кириллица или латиница, без апострофов)',
            'en': 'In English (Latin alphabet only, no apostrophes)',
        }.get(language, "o'zbek tilida apostrofsiz")

        # Ko'p so'z so'rash — AI ba'zilarini noto'g'ri qaytarsa ham yetishi uchun
        ask_count = max(count + 8, 25)

        prompt = f"""Sen krossvord uchun so'z va ta'rif tayyorlovchi yordamchisan.

Mavzu: "{topic}"

Vazifa: shu mavzu bo'yicha aniq {ask_count} ta so'z va ularning ta'riflarini (clue) yarat.

⚠️ KRITIK TALABLAR (qattiq rioya qiling):
1. Har so'z FAQAT lotin alifbosi (A-Z) — RAQAM, PROBEL, DEFIS, APOSTROF, '"', "'" YO'Q
2. APOSTROFLI so'zlarni QO'SHMA harflar bilan ALMASHTIRING:
   - "o'qituvchi" → "OQITUVCHI"
   - "ko'cha" → "KOCHA"
   - "g'oz" → "GOZ"
   - "ish'or" → "ISHOR"
3. So'z uzunligi: 4 dan 11 gacha harf (KESKIN)
4. Hammasi BIR-BIRIDAN FARQLI bo'lishi shart
5. Ta'rif (clue) — qisqa, aniq, 1 gap (5-12 so'z), {lang_label}
6. Ta'rif so'zning o'zini ishlatmasligi kerak (tarjima qilib boshqacha aytish)
7. Faqat JSON qaytar, boshqa hech narsa yozma

JSON formati (aynan shu, shu xil):
{{
  "words": [
    {{"word": "OQITUVCHI", "clue": "Maktabda dars beruvchi shaxs"}},
    {{"word": "MAKTAB", "clue": "Bilim olish joyi"}}
  ]
}}"""

        if attempt > 1:
            prompt += f"\n\n⚠️ AVVALGI URINISHDA so'zlarning ko'pi noto'g'ri formatda edi (apostrof, juda qisqa yoki uzun). Bu safar HAR BIR so'z qoidalarga to'liq mos kelishi shart!"

        response = await self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},
            temperature=0.6 if attempt == 1 else 0.4,
        )

        raw = response.choices[0].message.content
        try:
            data = json.loads(raw)
        except Exception as e:
            logger.warning(f"Krossvord JSON parse xato (attempt {attempt}): {e}")
            data = {"words": []}
        words = data.get("words", [])

        # Tozalash + validatsiya — har bir so'zni "saqlashga" urinamiz
        cleaned = []
        seen = set()
        for w in words:
            raw_word = str(w.get('word', '')).strip()
            # 1) Apostrof + tirelarni butunlay olib tashlash
            raw_word = re.sub(r"[’ʻʼˈ'`´\-\s]+", '', raw_word)
            # 2) FAQAT alifbo harflarini saqlash (lotin va kirill)
            word = re.sub(r'[^A-Za-zЀ-ӿ]', '', raw_word).upper()
            clue = str(w.get('clue', '')).strip()
            # Validatsiya
            if not (4 <= len(word) <= 11):
                continue
            if word in seen:
                continue
            if not clue:
                continue
            seen.add(word)
            cleaned.append({'word': word, 'clue': clue})

        logger.info(f"Krossvord AI [{attempt}-attempt]: {len(words)} ta so'z keldi, {len(cleaned)} ta tozalashdan o'tdi")

        # Agar kam bo'lsa va birinchi urinish bo'lsa, qayta so'raymiz
        if len(cleaned) < max(8, count - 5) and attempt < 2:
            logger.info(f"Krossvord uchun qayta so'rov (attempt 2)")
            extra = await self.generate_words(topic, count, language, attempt=2)
            # Birlashtirish (dublikatsiz)
            for entry in extra:
                if entry['word'] not in seen:
                    cleaned.append(entry)
                    seen.add(entry['word'])

        return cleaned


# ───────────────────────────────────────────────────────────────
# Asosiy interfeys
# ───────────────────────────────────────────────────────────────

async def generate_crossword(topic: str, openai_key: str, output_path: str,
                              word_count: int = 18, language: str = 'uz') -> bool:
    """End-to-end krossvord yaratish."""
    ai = CrosswordAIGenerator(openai_key)
    words_with_clues = await ai.generate_words(topic, word_count, language)
    if len(words_with_clues) < 5:
        logger.error(
            f"Krossvord uchun so'zlar yetarli emas: faqat {len(words_with_clues)} ta. "
            f"Mavzu juda tor yoki abstrakt bo'lishi mumkin: '{topic}'"
        )
        return False

    # Eng uzunidan qisqaga tartiblash (uzun so'z dastlab markazga, kichkina keyinroq joylanadi)
    words_with_clues.sort(key=lambda x: -len(x['word']))

    grid = CrosswordGrid(GRID_SIZE)
    placed_count = 0
    skipped = []
    for entry in words_with_clues:
        if grid.try_add(entry['word'], entry['clue']):
            placed_count += 1
        else:
            skipped.append(entry['word'])

    logger.info(f"Krossvord: {placed_count}/{len(words_with_clues)} so'z joylashtirildi. Skip: {skipped}")

    if placed_count < 5:
        logger.error("Juda kam so'z joylashtirildi, krossvord yaratilmadi")
        return False

    renderer = CrosswordDocxRenderer()
    return renderer.create_crossword(grid, topic, output_path)
