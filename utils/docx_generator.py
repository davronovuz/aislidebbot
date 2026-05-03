# utils/docx_generator.py
# DOCX FAYL YARATISH
# python-docx kutubxonasi bilan professional akademik hujjat yaratish
# GOST standarti bo'yicha (O'zbekiston akademik ishlari uchun)

from __future__ import annotations

import os
import re
import logging
from typing import Dict, List, Optional
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx kutubxonasi topilmadi!")


# Ish turi nomlari
WORK_TYPE_LABELS = {
    'mustaqil_ish': 'MUSTAQIL ISH',
    'referat': 'REFERAT',
    'kurs_ishi': 'KURS ISHI',
    'diplom_ishi': 'BITIRUV MALAKAVIY ISHI',
    'bitiruv_malakaviy_ishi': 'BITIRUV MALAKAVIY ISHI',
    'magistr_dissertatsiyasi': 'MAGISTR DISSERTATSIYASI',
    'laboratoriya_ishi': 'LABORATORIYA ISHI',
    'amaliy_ish': 'AMALIY ISH',
    'hisobot': 'HISOBOT',
    'ilmiy_maqola': 'ILMIY MAQOLA',
}


class DocxGenerator:
    """
    Professional DOCX hujjatlar yaratish.
    Mustaqil ish, referat, kurs ishi uchun.
    GOST standarti bo'yicha formatlash.
    """

    # Shrift nomi
    FONT_NAME = 'Times New Roman'

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kutubxonasi o'rnatilmagan. pip install python-docx")

        # Sahifa o'lchamlari (GOST)
        self.LEFT_MARGIN = Cm(3)
        self.RIGHT_MARGIN = Cm(1.5)
        self.TOP_MARGIN = Cm(2)
        self.BOTTOM_MARGIN = Cm(2)

        # Shrift o'lchamlari
        self.FONT_SIZE_BODY = Pt(14)
        self.FONT_SIZE_H1 = Pt(16)
        self.FONT_SIZE_H2 = Pt(14)
        self.FONT_SIZE_TITLE = Pt(18)
        self.FONT_SIZE_INSTITUTION = Pt(14)

        # Paragraf sozlamalari
        self.FIRST_LINE_INDENT = Cm(1.25)

    def create_course_work(
            self,
            content: Dict,
            output_path: str,
            work_type: str = 'mustaqil_ish'
    ) -> bool:
        """
        Mustaqil ish / Referat / Kurs ishi DOCX yaratish

        Args:
            content: Content dict
            output_path: Chiqish fayl yo'li
            work_type: Ish turi

        Returns:
            bool: Muvaffaqiyat
        """
        try:
            logger.info(f"DOCX yaratish boshlandi: {work_type}")

            doc = Document()

            # Sahifa sozlamalari (GOST + agar kurs ishi/diplom bo'lsa — sahifa ramkasi)
            self._setup_page(doc, work_type)

            # Stillar sozlash
            self._setup_styles(doc)

            # Sahifa raqamlash (2-sahifadan boshlab)
            self._add_page_numbers(doc)

            # Titul sahifa
            self._add_title_page(doc, content, work_type)

            # Annotatsiya sahifasi
            self._add_annotation_page(doc, content)

            # Mundarija
            self._add_table_of_contents(doc, content)

            # Kirish
            self._add_introduction(doc, content)

            # Asosiy boblar
            self._add_chapters(doc, content)

            # Xulosa
            self._add_conclusion(doc, content)

            # Tavsiyalar (agar bor bo'lsa)
            if content.get('recommendations'):
                self._add_recommendations(doc, content)

            # Adabiyotlar ro'yxati
            self._add_references(doc, content)

            # Ilovalar (agar bor bo'lsa)
            if content.get('appendix'):
                self._add_appendix(doc, content)

            # Saqlash
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            logger.info(f"DOCX saqlandi: {output_path}")

            return True

        except Exception as e:
            logger.error(f"DOCX yaratishda xato: {e}", exc_info=True)
            return False

    # ──────────────────────────────────────────────
    # Sahifa va stil sozlamalari
    # ──────────────────────────────────────────────

    # Page borders qo'yiladigan ish turlari (rasmiy hujjatlar)
    BORDERED_WORK_TYPES = {
        'kurs_ishi', 'diplom_ishi', 'bitiruv_malakaviy_ishi',
        'magistr_dissertatsiyasi', 'magistr_diss',
    }

    def _setup_page(self, doc: Document, work_type: str = ''):
        """Sahifa o'lchamlari va marginlarni sozlash (GOST).
        Agar work_type rasmiy hujjat bo'lsa — sahifa atrofiga ramka qo'yiladi.
        """
        section = doc.sections[0]
        section.left_margin = self.LEFT_MARGIN
        section.right_margin = self.RIGHT_MARGIN
        section.top_margin = self.TOP_MARGIN
        section.bottom_margin = self.BOTTOM_MARGIN

        if work_type in self.BORDERED_WORK_TYPES:
            self._add_page_borders(section)

    def _add_page_borders(self, section):
        """Section'ga sahifa atrofidagi ramka qo'shish (kurs ishi/diplom standarti)."""
        sectPr = section._sectPr
        # Eski pgBorders bo'lsa olib tashlash (idempotent)
        for old in sectPr.findall(qn('w:pgBorders')):
            sectPr.remove(old)
        pg_borders = parse_xml(
            f'<w:pgBorders {nsdecls("w")} w:offsetFrom="page" w:display="allPages">'
            f'<w:top w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="12" w:space="24" w:color="000000"/>'
            f'</w:pgBorders>'
        )
        sectPr.append(pg_borders)

    def _setup_styles(self, doc: Document):
        """Hujjat stillarini sozlash"""
        # Normal stil
        style = doc.styles['Normal']
        font = style.font
        font.name = self.FONT_NAME
        font.size = self.FONT_SIZE_BODY
        font.color.rgb = RGBColor(0, 0, 0)

        para_format = style.paragraph_format
        para_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para_format.first_line_indent = self.FIRST_LINE_INDENT
        para_format.space_after = Pt(0)
        para_format.space_before = Pt(0)
        para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # GOST: ikki tomonlama tekislash
        para_format.widow_control = True  # Yetim qatorlarni oldini olish
        para_format.keep_together = False

        # Heading 1 - Bob sarlavhasi
        if 'Heading 1' in doc.styles:
            h1 = doc.styles['Heading 1']
            h1.font.name = self.FONT_NAME
            h1.font.size = self.FONT_SIZE_H1
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(0, 0, 0)
            h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h1.paragraph_format.space_before = Pt(0)
            h1.paragraph_format.space_after = Pt(12)
            h1.paragraph_format.first_line_indent = Cm(0)
            h1.paragraph_format.keep_with_next = True  # Sarlavha keyingi paragrafdan ajralmasin
            h1.paragraph_format.page_break_before = False

        # Heading 2 - Bo'lim sarlavhasi
        if 'Heading 2' in doc.styles:
            h2 = doc.styles['Heading 2']
            h2.font.name = self.FONT_NAME
            h2.font.size = self.FONT_SIZE_H2
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(0, 0, 0)
            h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h2.paragraph_format.space_before = Pt(12)
            h2.paragraph_format.space_after = Pt(6)
            h2.paragraph_format.first_line_indent = Cm(0)
            h2.paragraph_format.keep_with_next = True  # Sarlavha keyingi paragrafdan ajralmasin

        # List Bullet stili yaratish yoki sozlash
        try:
            list_style = doc.styles.add_style('BulletList', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            list_style = doc.styles['BulletList']
        list_style.font.name = self.FONT_NAME
        list_style.font.size = self.FONT_SIZE_BODY
        list_style.paragraph_format.first_line_indent = Cm(0)
        list_style.paragraph_format.left_indent = Cm(1.25)
        list_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        list_style.paragraph_format.space_after = Pt(0)

        # Numbered list stili
        try:
            num_style = doc.styles.add_style('NumberedList', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            num_style = doc.styles['NumberedList']
        num_style.font.name = self.FONT_NAME
        num_style.font.size = self.FONT_SIZE_BODY
        num_style.paragraph_format.first_line_indent = Cm(0)
        num_style.paragraph_format.left_indent = Cm(1.25)
        num_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        num_style.paragraph_format.space_after = Pt(0)

    def _add_page_numbers(self, doc: Document):
        """
        Sahifa raqamlarini footer ga qo'shish.
        Titul sahifada raqam ko'rinmaydi (2-sahifadan boshlab).
        """
        section = doc.sections[0]
        # Birinchi sahifa boshqacha (raqamsiz)
        section.different_first_page_header_footer = True

        # Asosiy footer - sahifa raqami markazda
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.first_line_indent = Cm(0)

        # PAGE field qo'shish
        run = footer_para.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run2 = footer_para.add_run()
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr_text)

        run3 = footer_para.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)

        # Shrift sozlash
        for r in [run, run2, run3]:
            r.font.name = self.FONT_NAME
            r.font.size = Pt(12)

        # Birinchi sahifa footer - bo'sh
        first_footer = section.first_page_footer
        first_footer.is_linked_to_previous = False
        if first_footer.paragraphs:
            first_footer.paragraphs[0].text = ''

    # ──────────────────────────────────────────────
    # Matn formatlash yordamchilari
    # ──────────────────────────────────────────────

    def _parse_and_add_formatted_text(self, paragraph, text: str):
        """
        Matnni parse qilib, **bold** belgilarini formatlash.
        **matn** -> qalin qilib yoziladi.
        """
        # **text** patternini topish
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.name = self.FONT_NAME
                run.font.size = self.FONT_SIZE_BODY
            else:
                if part:
                    run = paragraph.add_run(part)
                    run.font.name = self.FONT_NAME
                    run.font.size = self.FONT_SIZE_BODY

    def _add_formatted_content(self, doc: Document, text: str):
        """
        Matnni paragraflar, ro'yxatlar, va formatlangan elementlarga ajratib qo'shish.

        Qo'llab-quvvatlanadigan formatlar:
        - Oddiy paragraflar (bo'sh qator bilan ajratilgan)
        - Bullet ro'yxatlar (- yoki * bilan boshlanuvchi qatorlar)
        - Raqamli ro'yxatlar (1. 2. 3. bilan boshlanuvchi qatorlar)
        - **qalin matn** formatlash
        """
        if not text:
            return

        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()

            if not line:
                # Bo'sh qator - o'tkazib yuborish
                i += 1
                continue

            # Bullet ro'yxat
            if re.match(r'^[-\u2022\u2023\u25E6\u2043\u2219]\s+', line):
                content = re.sub(r'^[-\u2022\u2023\u25E6\u2043\u2219]\s+', '', line)
                p = doc.add_paragraph()
                p.style = doc.styles['BulletList']
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(1.25)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                bullet_run = p.add_run('\u2022 ')
                bullet_run.font.name = self.FONT_NAME
                bullet_run.font.size = self.FONT_SIZE_BODY

                self._parse_and_add_formatted_text(p, content)
                i += 1
                continue

            # Raqamli ro'yxat (1. 2. 3. ...)
            num_match = re.match(r'^(\d+)[.)]\s+', line)
            if num_match:
                num = num_match.group(1)
                content = re.sub(r'^\d+[.)]\s+', '', line)
                p = doc.add_paragraph()
                p.style = doc.styles['NumberedList']
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.left_indent = Cm(1.25)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                num_run = p.add_run(f'{num}. ')
                num_run.font.name = self.FONT_NAME
                num_run.font.size = self.FONT_SIZE_BODY

                self._parse_and_add_formatted_text(p, content)
                i += 1
                continue

            # Oddiy paragraf
            # Ketma-ket qatorlarni bitta paragrafga birlashtirish
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if not next_line:
                    i += 1
                    break
                # Agar keyingi qator ro'yxat elementi bo'lsa, to'xtatish
                if re.match(r'^[-\u2022\u2023\u25E6\u2043\u2219]\s+', next_line):
                    break
                if re.match(r'^\d+[.)]\s+', next_line):
                    break
                para_lines.append(next_line)
                i += 1

            full_text = ' '.join(para_lines)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            self._parse_and_add_formatted_text(p, full_text)

    def _add_section_heading(self, doc: Document, text: str, level: int = 1,
                              add_page_break: bool = False):
        """
        Sarlavha qo'shish.
        level=1: Bob sarlavhasi (markazlashtirilgan, 16pt, bold + pastki chiziq)
        level=2: Bo'lim sarlavhasi (chapdan, 14pt, bold)
        """
        if add_page_break:
            doc.add_page_break()

        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0) if add_page_break else Pt(14)
        p.paragraph_format.space_after = Pt(14) if level == 1 else Pt(8)
        p.paragraph_format.keep_with_next = True

        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.upper())
            run.bold = True
            run.font.size = self.FONT_SIZE_H1
            # Bob sarlavhasi tagiga chiziq qo'shish (paragraph bottom border)
            pPr = p._p.get_or_add_pPr()
            border_xml = (
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="6" w:space="6" w:color="000000"/>'
                f'</w:pBdr>'
            )
            pPr.append(parse_xml(border_xml))
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.size = self.FONT_SIZE_H2

        run.font.name = self.FONT_NAME
        run.font.color.rgb = RGBColor(0, 0, 0)

    # ──────────────────────────────────────────────
    # Hujjat bo'limlari
    # ──────────────────────────────────────────────

    def _add_centered_line(self, doc: Document, text: str, size: int = 14, bold: bool = False, space_after: int = 0):
        """Markazlashtirilgan qator qo'shish"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = self.FONT_NAME
        return p

    def _add_empty_lines(self, doc: Document, count: int):
        """Bo'sh qatorlar qo'shish (compact, titul sahifaga sig'ishi uchun)"""
        for _ in range(count):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run('')
            run.font.size = Pt(12)
            run.font.name = self.FONT_NAME

    def _add_right_line(self, doc: Document, label: str, value: str, size: int = 14):
        """O'ngga tekislangan label: value qator"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        # Label (oddiy)
        run = p.add_run(f'{label} ')
        run.font.size = Pt(size)
        run.font.name = self.FONT_NAME
        # Value (qalin)
        run = p.add_run(value if value else '_______________')
        run.bold = bool(value)
        run.font.size = Pt(size)
        run.font.name = self.FONT_NAME

    def _add_title_page(self, doc: Document, content: Dict, work_type: str):
        """Titul sahifa — ish turi bo'yicha tegishli shablon tanlanadi."""
        if work_type in self.BORDERED_WORK_TYPES:
            self._add_official_title_page(doc, content, work_type)
        else:
            self._add_simple_title_page(doc, content, work_type)

    def _add_official_title_page(self, doc: Document, content: Dict, work_type: str):
        """
        Rasmiy titul sahifa — kurs ishi va diplom uchun (O'zbekiston standarti).
        Sahifa atrofida ramka, ulkan ish turi nomi (44pt), markazlashgan.
        Foydalanuvchi yuborgan namuna asosida.
        """
        author_info = content.get('author_info', {})
        subject = content.get('subject', '')
        title = content.get('title', 'MAVZU')

        # ─── YUQORI: Vazirlik (rasmda 2 qator alohida)
        self._add_centered_line(doc, "O'ZBEKISTON RESPUBLIKASI",
            size=13, bold=True, space_after=4)
        self._add_centered_line(doc,
            "OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI",
            size=13, bold=True, space_after=10)

        # Universitet (ostiga chiziq)
        institution = author_info.get('institution', '')
        self._add_centered_with_underline(doc, institution.upper() if institution else '',
            size=12, bold=True)

        # Fakultet (ostiga chiziq)
        faculty = author_info.get('faculty', '')
        self._add_centered_with_underline(doc,
            f"{faculty} fakulteti" if faculty else '', size=12)

        # Kafedra
        department = author_info.get('department', '')
        if department:
            self._add_centered_with_underline(doc, department, size=12)
        elif subject:
            self._add_centered_with_underline(doc, f'"{subject}" kafedrasi', size=12)

        # Kurs / guruh (rasmdagi __-KURS ___-GURUH TALABASI)
        student_group = author_info.get('student_group', '')
        course_text = (f"{student_group} guruh talabasi" if student_group
                       else "____ kurs ____ guruh talabasi")
        self._add_centered_line(doc, course_text, size=12, space_after=8)

        # Talaba ismi NING (rasmdagi pattern)
        student = author_info.get('student_name', '')
        if student:
            self._add_centered_with_underline(doc, student.upper(), size=13, bold=True, suffix=' NING')
        else:
            self._add_centered_with_underline(doc, '', size=13, bold=True, suffix=' NING')

        self._add_empty_lines(doc, 1)

        # Mavzu: "_____ mavzusidagi"
        self._add_centered_with_underline(doc, title, size=12)
        self._add_centered_line(doc, "mavzusidagi", size=12, space_after=12)

        # ─── ASOSIY: ULKAN ISH TURI (rasmda 50pt+ ko'rinadi)
        work_label = WORK_TYPE_LABELS.get(work_type, 'KURS ISHI')
        self._add_centered_line(doc, work_label, size=44, bold=True, space_after=4)

        # ─── PASTKI: rahbar (ixtiyoriy)
        self._add_empty_lines(doc, 6)

        teacher = author_info.get('teacher_name', '')
        teacher_rank = author_info.get('teacher_rank', '')
        if teacher:
            value = f'{teacher_rank} {teacher}' if teacher_rank else teacher
            self._add_right_line(doc, 'Ilmiy rahbar:', value, size=12)

        # Pastda Toshkent – yil
        self._add_empty_lines(doc, 4)
        self._add_centered_line(doc,
            f"Toshkent {chr(0x2013)} {datetime.now().year}",
            size=13, bold=True)

        doc.add_page_break()

    def _add_centered_with_underline(self, doc: Document, text: str, size: int = 12,
                                      bold: bool = False, suffix: str = ''):
        """
        Markazlashgan matn + ostida chiziq (rasmdagi pattern).
        Agar matn bo'sh bo'lsa, faqat chiziq chiqadi (qo'lda to'ldirish uchun).
        """
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.line_spacing = 1.2

        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = self.FONT_NAME
            run.underline = True
        else:
            run = p.add_run('_' * 60)
            run.font.size = Pt(size)
            run.font.name = self.FONT_NAME

        if suffix:
            run2 = p.add_run(suffix)
            run2.bold = bold
            run2.font.size = Pt(size)
            run2.font.name = self.FONT_NAME

    def _add_simple_title_page(self, doc: Document, content: Dict, work_type: str):
        """
        Sodda titul — mustaqil ish, referat va h.k. (ramkasiz, kompakt).
        """
        author_info = content.get('author_info', {})
        subject = content.get('subject', '')
        title = content.get('title', 'MAVZU')

        # YUQORI: Vazirlik (14pt bold)
        self._add_centered_line(doc,
            "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI",
            size=14, bold=True, space_after=8)

        # Universitet nomi (16pt bold — asosiy diqqat)
        institution = author_info.get('institution', '')
        if institution:
            self._add_centered_line(doc, institution.upper(), size=16, bold=True, space_after=6)

        # Fakultet
        faculty = author_info.get('faculty', '')
        if faculty:
            self._add_centered_line(doc, f"{faculty} fakulteti", size=14, bold=False, space_after=4)

        # Kafedra (fan nomi asosida)
        department = author_info.get('department', '')
        if department:
            self._add_centered_line(doc, department, size=14, bold=False, space_after=0)
        elif subject:
            self._add_centered_line(doc, f'"{subject}" kafedrasi', size=14, bold=False, space_after=0)

        # Yuqori qismni ajratuvchi gorizontal chiziq
        self._add_horizontal_separator(doc, space_before=6, space_after=6)

        # ─── O'RTA QISM: Ish turi + Mavzu (markaziy diqqat) ───

        self._add_empty_lines(doc, 2)

        # Ish turi (22pt bold — yirik)
        work_label = WORK_TYPE_LABELS.get(work_type, 'MUSTAQIL ISH')
        self._add_centered_line(doc, work_label, size=22, bold=True, space_after=10)

        # "Mavzu:" sarlavhasi (16pt bold)
        self._add_centered_line(doc, 'Mavzu:', size=16, bold=True, space_after=6)

        # Mavzu nomi — KATTA (24pt bold) + ramka
        title = content.get('title', 'MAVZU')
        self._add_bordered_title(doc, f'«{title.upper()}»', size=24)

        # Fan nomi (subtitle)
        if subject:
            self._add_centered_line(doc, f'{subject} fanidan', size=14, bold=False, space_after=0)

        # ─── PASTKI-O'RTA QISM: Bajardi / Tekshirdi ───

        self._add_empty_lines(doc, 2)

        # Bajardi bloki
        student = author_info.get('student_name', '')
        student_group = author_info.get('student_group', '')
        self._add_right_line(doc, 'Bajardi:', student)
        if student_group:
            self._add_right_line(doc, 'Guruh:', student_group)

        # Bo'sh qator
        self._add_empty_lines(doc, 1)

        # Tekshirdi bloki
        teacher = author_info.get('teacher_name', '')
        teacher_rank = author_info.get('teacher_rank', '')
        if teacher_rank and teacher:
            self._add_right_line(doc, 'Tekshirdi:', f'{teacher_rank} {teacher}')
        else:
            self._add_right_line(doc, 'Tekshirdi:', teacher)

        # ─── ENG PASTKI QISM: Shahar – Yil ───

        self._add_empty_lines(doc, 1)
        self._add_centered_line(doc, f'Toshkent \u2013 {datetime.now().year}', size=14, bold=True)

        # Yangi sahifa
        doc.add_page_break()

    def _add_horizontal_separator(self, doc: Document, space_before: int = 6, space_after: int = 6):
        """Gorizontal ajratuvchi chiziq (paragraph bottom border orqali)."""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        pPr = p._p.get_or_add_pPr()
        border_xml = (
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(parse_xml(border_xml))

    def _add_bordered_title(self, doc: Document, text: str, size: int = 24):
        """Mavzu nomini ramka ichida, markazlashtirilgan, katta shrift bilan qo'shish (titul markaziy diqqat)."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.right_indent = Cm(1.5)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.2

        pPr = p._p.get_or_add_pPr()
        border_xml = (
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="6" w:color="000000"/>'
            f'<w:left w:val="single" w:sz="12" w:space="8" w:color="000000"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="6" w:color="000000"/>'
            f'<w:right w:val="single" w:sz="12" w:space="8" w:color="000000"/>'
            f'</w:pBdr>'
        )
        pPr.append(parse_xml(border_xml))

        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = self.FONT_NAME
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_annotation_page(self, doc: Document, content: Dict):
        """
        Annotatsiya sahifasi.
        Abstract va kalit so'zlar.
        """
        abstract = content.get('abstract', '')
        keywords = content.get('keywords', [])

        if not abstract and not keywords:
            return

        # Sarlavha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('ANNOTATSIYA')
        run.bold = True
        run.font.size = self.FONT_SIZE_H1
        run.font.name = self.FONT_NAME

        # Abstract matn
        if abstract:
            self._add_formatted_content(doc, abstract)

        # Kalit so'zlar
        if keywords:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run('Kalit so\'zlar: ')
            run.bold = True
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE_BODY

            keywords_text = ', '.join(keywords) + '.'
            run = p.add_run(keywords_text)
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE_BODY

        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document, content: Dict):
        """
        Mundarija sahifasi.
        Nuqtali chiziq va o'ngga tekislangan sahifa raqamlari bilan.
        """
        # Sarlavha
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run('MUNDARIJA')
        run.bold = True
        run.font.size = self.FONT_SIZE_H1
        run.font.name = self.FONT_NAME

        toc = content.get('table_of_contents', [])

        # Sahifa kengligi hisoblab tab stop pozitsiyasi
        # A4: 21cm - 3cm left - 1.5cm right = 16.5cm
        tab_stop_position = Cm(16.5)

        for item in toc:
            title = item.get('title', '')
            page = item.get('page', '')

            # 1-darajali band (bob, kirish, xulosa, adabiyotlar) — bold
            # 2-darajali band (1.1., 1.2. ...) — oddiy + chap chetidan biroz ichkari
            stripped = title.lstrip()
            is_subsection = bool(re.match(r'^\d+\.\d+', stripped))
            is_top_level = (not is_subsection) and (
                stripped.isupper()
                or re.match(r'^\d+[-.]?\s*BOB', stripped, re.IGNORECASE)
                or stripped.upper().startswith(('KIRISH', 'XULOSA', 'MUNDARIJA',
                                                'ANNOTATSIYA', 'TAVSIYALAR',
                                                'FOYDALANILGAN', 'ILOVALAR'))
            )

            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0.6) if is_subsection else Cm(0)
            p.paragraph_format.space_after = Pt(3) if is_top_level else Pt(2)
            p.paragraph_format.space_before = Pt(6) if is_top_level else Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # Tab stop qo'shish - nuqtali chiziq bilan o'ng tomonga
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                tab_stop_position,
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS
            )

            # Sarlavha matni
            run = p.add_run(title)
            run.font.size = self.FONT_SIZE_BODY
            run.font.name = self.FONT_NAME
            if is_top_level:
                run.bold = True

            # Tab va sahifa raqami
            if page is not None and page != '':
                run = p.add_run(f'\t{page}')
                run.font.size = self.FONT_SIZE_BODY
                run.font.name = self.FONT_NAME
                if is_top_level:
                    run.bold = True

        doc.add_page_break()

    def _add_introduction(self, doc: Document, content: Dict):
        """Kirish bo'limi"""
        intro = content.get('introduction', {})
        if not intro:
            return

        title = intro.get('title', 'KIRISH')
        text = intro.get('content', '')

        self._add_section_heading(doc, title, level=1, add_page_break=False)
        self._add_formatted_content(doc, text)
        # Keyingi bob `add_page_break=True` bilan boshlanadi — bu yerda qo'shimcha break shart emas

    def _add_chapters(self, doc: Document, content: Dict):
        """Asosiy boblarni qo'shish (har bir bob yangi sahifadan)"""
        chapters = content.get('chapters', [])

        for chapter in chapters:
            number = chapter.get('number', 1)
            title = chapter.get('title', '')

            # Bob sarlavhasi (yangi sahifadan)
            chapter_title = f'{number}-BOB. {title.upper()}'
            self._add_section_heading(doc, chapter_title, level=1, add_page_break=True)

            # Bo'limlar (bir bob ichida sahifa uzilishi yo'q)
            sections = chapter.get('sections', [])
            for section in sections:
                sec_number = section.get('number', '')
                sec_title = section.get('title', '')
                sec_content = section.get('content', '')

                # Bo'lim sarlavhasi
                section_title = f'{sec_number}. {sec_title}'
                self._add_section_heading(doc, section_title, level=2, add_page_break=False)

                # Bo'lim matni
                self._add_formatted_content(doc, sec_content)

    def _add_conclusion(self, doc: Document, content: Dict):
        """Xulosa bo'limi"""
        conclusion = content.get('conclusion', {})
        if not conclusion:
            return

        title = conclusion.get('title', 'XULOSA')
        text = conclusion.get('content', '')

        self._add_section_heading(doc, title, level=1, add_page_break=True)
        self._add_formatted_content(doc, text)

    def _add_recommendations(self, doc: Document, content: Dict):
        """Tavsiyalar bo'limi"""
        recommendations = content.get('recommendations', [])
        if not recommendations:
            return

        self._add_section_heading(doc, 'TAVSIYALAR', level=1, add_page_break=True)

        for i, rec in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            run = p.add_run(f'{i}. ')
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE_BODY

            self._parse_and_add_formatted_text(p, rec)

    def _add_references(self, doc: Document, content: Dict):
        """Adabiyotlar ro'yxati"""
        references = content.get('references', [])
        if not references:
            return

        self._add_section_heading(
            doc, "FOYDALANILGAN ADABIYOTLAR RO'YXATI",
            level=1, add_page_break=True
        )

        for i, ref in enumerate(references, 1):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.hanging_indent = Cm(1.25)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p.paragraph_format.space_after = Pt(2)

            # Agar raqam bilan boshlanmasa, raqam qo'shish
            if not re.match(r'^\d+[.)]\s', ref):
                run = p.add_run(f'{i}. {ref}')
            else:
                run = p.add_run(ref)
            run.font.size = self.FONT_SIZE_BODY
            run.font.name = self.FONT_NAME

    def _add_appendix(self, doc: Document, content: Dict):
        """Ilovalar bo'limi"""
        appendix = content.get('appendix')
        if not appendix:
            return

        self._add_section_heading(doc, 'ILOVALAR', level=1, add_page_break=True)

        if isinstance(appendix, str):
            self._add_formatted_content(doc, appendix)
        elif isinstance(appendix, list):
            for idx, item in enumerate(appendix):
                if isinstance(item, dict):
                    item_title = item.get('title', f'Ilova {idx + 1}')
                    item_content = item.get('content', '')
                    if idx > 0:
                        doc.add_page_break()
                    self._add_section_heading(doc, item_title, level=2, add_page_break=False)
                    self._add_formatted_content(doc, str(item_content))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = self.FIRST_LINE_INDENT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                    self._parse_and_add_formatted_text(p, str(item))


# Helper function - tashqi interface sifatida ishlatiladi
def create_docx_from_content(content: Dict, output_path: str, work_type: str = 'mustaqil_ish') -> bool:
    """DOCX yaratish helper"""
    try:
        generator = DocxGenerator()
        return generator.create_course_work(content, output_path, work_type)
    except Exception as e:
        logger.error(f"DOCX helper xato: {e}", exc_info=True)
        return False
